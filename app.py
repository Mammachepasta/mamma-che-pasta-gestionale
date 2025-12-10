import sqlite3
from datetime import datetime, date, timedelta
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import csv
import io
from docx import Document

app = Flask(__name__)
app.secret_key = "chiave-super-segreta"

DB_PATH = "gestionale.db"


# ---------------------- DB UTILS ----------------------


def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db_connection()
    cur = conn.cursor()

    # CLIENTI
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codice TEXT,
            nome TEXT NOT NULL UNIQUE
        )
        """
    )

    # PRODOTTI
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS prodotti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codice TEXT,
            nome TEXT NOT NULL UNIQUE,
            kg_per_vaschetta REAL NOT NULL,
            giacenza_iniziale_vaschette REAL NOT NULL DEFAULT 0
        )
        """
    )

    # ORDINI (testata)
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS ordini (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data TEXT NOT NULL,
            cliente_id INTEGER NOT NULL,
            FOREIGN KEY (cliente_id) REFERENCES clienti(id)
        )
        """
    )

    # RIGHE ORDINE (dettaglio)
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS righe_ordine (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ordine_id INTEGER NOT NULL,
            prodotto_id INTEGER NOT NULL,
            qta_inserita REAL NOT NULL,
            tipo_qta TEXT NOT NULL,
            FOREIGN KEY (ordine_id) REFERENCES ordini(id),
            FOREIGN KEY (prodotto_id) REFERENCES prodotti(id)
        )
        """
    )

    # PRODUZIONE
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS produzione (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            data TEXT NOT NULL,
            prodotto_id INTEGER NOT NULL,
            vaschette_prodotte REAL NOT NULL,
            FOREIGN KEY (prodotto_id) REFERENCES prodotti(id)
        )
        """
    )

    conn.commit()
    conn.close()


# ---------------------- FUNZIONI LOGICHE ----------------------


def calcola_magazzino():
    """
    Calcola la giacenza per ogni prodotto, in vaschette e in kg.
    """
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("SELECT * FROM prodotti ORDER BY nome")
    prodotti = cur.fetchall()

    magazzino = []

    for p in prodotti:
        pid = p["id"]

        # produzione totale (vaschette)
        cur.execute(
            "SELECT COALESCE(SUM(vaschette_prodotte), 0) FROM produzione WHERE prodotto_id = ?",
            (pid,),
        )
        prodotte = cur.fetchone()[0] or 0

        # ordini totali (convertiti in vaschette)
        cur.execute(
            "SELECT qta_inserita, tipo_qta FROM righe_ordine WHERE prodotto_id = ?",
            (pid,),
        )
        righe = cur.fetchall()

        ordinate_v = 0
        for r in righe:
            if r["tipo_qta"] == "v":
                ordinate_v += r["qta_inserita"]
            else:
                if p["kg_per_vaschetta"] > 0:
                    ordinate_v += r["qta_inserita"] / p["kg_per_vaschetta"]

        giac_finale_v = p["giacenza_iniziale_vaschette"] + prodotte - ordinate_v
        giac_finale_kg = giac_finale_v * p["kg_per_vaschetta"]

        magazzino.append(
            {
                "id": pid,
                "codice": p["codice"],
                "nome": p["nome"],
                "kg_per_vaschetta": p["kg_per_vaschetta"],
                "giacenza_iniziale_v": p["giacenza_iniziale_vaschette"],
                "prodotte_v": prodotte,
                "ordinate_v": ordinate_v,
                "giacenza_finale_v": giac_finale_v,
                "giacenza_finale_kg": giac_finale_kg,
            }
        )

    conn.close()
    return magazzino


# ---------------------- ROUTE PRINCIPALE ----------------------


@app.route("/")
def index():
    magazzino = calcola_magazzino()
    return render_template("index.html", magazzino=magazzino)


# ---------------------- CLIENTI ----------------------


@app.route("/clienti", methods=["GET", "POST"])
def clienti():
    conn = get_db_connection()
    cur = conn.cursor()

    if request.method == "POST":
        codice = request.form.get("codice", "").strip() or None
        nome = request.form.get("nome", "").strip()

        if not nome:
            flash("Il nome è obbligatorio", "danger")
            return redirect(url_for("clienti"))

        try:
            cur.execute(
                "INSERT INTO clienti (codice, nome) VALUES (?, ?)",
                (codice, nome),
            )
            conn.commit()
            flash("Cliente aggiunto", "success")
        except sqlite3.IntegrityError:
            flash("Cliente già esistente", "danger")

        conn.close()
        return redirect(url_for("clienti"))

    cur.execute("SELECT * FROM clienti ORDER BY nome")
    clienti_rows = cur.fetchall()
    conn.close()
    return render_template("clienti.html", clienti=clienti_rows)


@app.route("/clienti/<int:id>/elimina", methods=["POST"])
def elimina_cliente(id):
    conn = get_db_connection()
    cur = conn.cursor()

    # controllo ordini collegati
    cur.execute("SELECT COUNT(*) FROM ordini WHERE cliente_id = ?", (id,))
    if cur.fetchone()[0] > 0:
        flash("Impossibile eliminare: cliente con ordini esistenti.", "danger")
        conn.close()
        return redirect(url_for("clienti"))

    cur.execute("DELETE FROM clienti WHERE id = ?", (id,))
    conn.commit()
    conn.close()
    flash("Cliente eliminato.", "info")
    return redirect(url_for("clienti"))


# ---------------------- PRODOTTI ----------------------


@app.route("/prodotti", methods=["GET", "POST"])
def prodotti():
    conn = get_db_connection()
    cur = conn.cursor()

    if request.method == "POST":
        codice = request.form.get("codice", "").strip() or None
        nome = request.form.get("nome", "").strip()
        kg_v = request.form.get("kg_per_vaschetta", "").replace(",", ".")
        giac_iniz = request.form.get("giacenza_iniziale_v", "").replace(",", ".")

        try:
            kg_v = float(kg_v)
            giac_iniz = float(giac_iniz or 0)
        except ValueError:
            flash("Numeri non validi per kg/vaschetta o giacenza iniziale.", "danger")
            return redirect(url_for("prodotti"))

        if not nome:
            flash("Il nome del prodotto è obbligatorio.", "danger")
            return redirect(url_for("prodotti"))

        try:
            cur.execute(
                """
                INSERT INTO prodotti (codice, nome, kg_per_vaschetta, giacenza_iniziale_vaschette)
                VALUES (?, ?, ?, ?)
                """,
                (codice, nome, kg_v, giac_iniz),
            )
            conn.commit()
            flash("Prodotto aggiunto.", "success")
        except sqlite3.IntegrityError:
            flash("Prodotto già esistente.", "danger")

        conn.close()
        return redirect(url_for("prodotti"))

    cur.execute("SELECT * FROM prodotti ORDER BY nome")
    prodotti_rows = cur.fetchall()
    conn.close()
    return render_template("prodotti.html", prodotti=prodotti_rows)


@app.route("/prodotti/<int:id>/elimina", methods=["POST"])
def elimina_prodotto(id):
    conn = get_db_connection()
    cur = conn.cursor()

    # controllo movimenti collegati
    cur.execute("SELECT COUNT(*) FROM righe_ordine WHERE prodotto_id = ?", (id,))
    cnt1 = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM produzione WHERE prodotto_id = ?", (id,))
    cnt2 = cur.fetchone()[0]

    if cnt1 + cnt2 > 0:
        flash("Impossibile eliminare: il prodotto ha movimenti registrati.", "danger")
    else:
        cur.execute("DELETE FROM prodotti WHERE id = ?", (id,))
        conn.commit()
        flash("Prodotto eliminato.", "info")

    conn.close()
    return redirect(url_for("prodotti"))


# ---------------------- ORDINI ----------------------

@app.route("/ordini")
def lista_ordini():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT 
            o.id AS id,
            o.data,
            c.nome AS cliente_nome,
            c.codice AS cliente_codice,
            COUNT(ro.id) AS num_righe,
            SUM(
                CASE 
                    WHEN ro.tipo_qta = 'kg' THEN ro.qta_inserita
                    ELSE ro.qta_inserita * p.kg_per_vaschetta
                END
            ) AS kg_totali
        FROM ordini o
        JOIN clienti c ON c.id = o.cliente_id
        JOIN righe_ordine ro ON ro.ordine_id = o.id
        JOIN prodotti p ON p.id = ro.prodotto_id
        GROUP BY o.id, o.data, c.nome, c.codice
        ORDER BY o.data DESC, o.id DESC
    """)
    ordini = cur.fetchall()
    conn.close()
    return render_template("ordini.html", ordini=ordini)


@app.route("/ordini/nuovo", methods=["GET", "POST"])
def nuovo_ordine():
    conn = get_db_connection()
    cur = conn.cursor()

    if request.method == "POST":
        data_str = request.form.get("data") or datetime.today().strftime("%Y-%m-%d")
        cliente_id = request.form.get("cliente_id")

        if not cliente_id:
            flash("Seleziona un cliente.", "danger")
            conn.close()
            return redirect(url_for("nuovo_ordine"))

        # crea testata ordine
        cur.execute(
            "INSERT INTO ordini (data, cliente_id) VALUES (?, ?)",
            (data_str, cliente_id),
        )
        ordine_id = cur.lastrowid

        # righe ordine (10 righe massimo fisse, semplici)
        righe_ok = 0
        for index in range(10):
            prod_id = request.form.get(f"prodotto_{index}")
            qta_str = request.form.get(f"qta_{index}", "").replace(",", ".")
            tipo = request.form.get(f"tipo_{index}")

            if not prod_id and not qta_str:
                continue

            if not prod_id or not qta_str:
                continue

            try:
                qta = float(qta_str)
            except ValueError:
                qta = 0

            if qta <= 0:
                continue

            if tipo not in ("kg", "v"):
                continue

            cur.execute(
                "INSERT INTO righe_ordine (ordine_id, prodotto_id, qta_inserita, tipo_qta) "
                "VALUES (?, ?, ?, ?)",
                (ordine_id, prod_id, qta, tipo),
            )
            righe_ok += 1

        if righe_ok == 0:
            # nessuna riga valida -> cancello l’ordine
            cur.execute("DELETE FROM ordini WHERE id = ?", (ordine_id,))
            conn.commit()
            conn.close()
            flash("Nessuna riga valida inserita.", "danger")
            return redirect(url_for("nuovo_ordine"))

        conn.commit()
        conn.close()
        flash("Ordine salvato correttamente.", "success")
        return redirect(url_for("lista_ordini"))

    # GET: form vuoto
    cur.execute("SELECT * FROM clienti ORDER BY nome")
    clienti = cur.fetchall()
    cur.execute("SELECT * FROM prodotti ORDER BY nome")
    prodotti = cur.fetchall()
    conn.close()
    return render_template("nuovo_ordine.html", clienti=clienti, prodotti=prodotti)


@app.route("/ordini/<int:ordine_id>/dettaglio")
def dettaglio_ordine(ordine_id):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT o.id AS id,
               o.data,
               c.nome AS cliente_nome,
               c.codice AS cliente_codice
        FROM ordini o
        JOIN clienti c ON c.id = o.cliente_id
        WHERE o.id = ?
    """, (ordine_id,))
    ordine = cur.fetchone()

    if ordine is None:
        conn.close()
        flash("Ordine non trovato.", "danger")
        return redirect(url_for("lista_ordini"))

    cur.execute("""
        SELECT ro.id,
               ro.qta_inserita,
               ro.tipo_qta,
               p.nome AS prodotto_nome,
               p.codice AS prodotto_codice,
               p.kg_per_vaschetta
        FROM righe_ordine ro
        JOIN prodotti p ON p.id = ro.prodotto_id
        WHERE ro.ordine_id = ?
        ORDER BY p.nome
    """, (ordine_id,))
    righe = cur.fetchall()
    conn.close()

    righe_calc = []
    tot_kg = 0.0
    tot_v = 0.0

    for r in righe:
        q = r["qta_inserita"]
        t = r["tipo_qta"]
        kg_v = r["kg_per_vaschetta"] or 0

        if t == "kg":
            kg = q
            v = q / kg_v if kg_v else 0
        else:  # vaschette
            v = q
            kg = q * kg_v

        tot_kg += kg
        tot_v += v

        righe_calc.append({
            "prodotto": r["prodotto_nome"],
            "codice": r["prodotto_codice"],
            "kg": kg,
            "vaschette": v,
        })

    return render_template(
        "dettaglio_ordine.html",
        ordine=ordine,
        righe=righe_calc,
        tot_kg=tot_kg,
        tot_vaschette=tot_v,
    )


@app.route("/ordini/<int:ordine_id>/elimina", methods=["POST"])
def elimina_ordine(ordine_id):
    conn = get_db_connection()
    cur = conn.cursor()
    # prima elimino righe
    cur.execute("DELETE FROM righe_ordine WHERE ordine_id = ?", (ordine_id,))
    # poi testata
    cur.execute("DELETE FROM ordini WHERE id = ?", (ordine_id,))
    conn.commit()
    conn.close()
    flash("Ordine eliminato.", "info")
    return redirect(url_for("lista_ordini"))



# ---------------------- PRODUZIONE ----------------------


@app.route("/produzione", methods=["GET", "POST"])
def produzione():
    conn = get_db_connection()
    cur = conn.cursor()

    if request.method == "POST":
        data = request.form.get("data") or datetime.today().strftime("%Y-%m-%d")
        prodotto_id = request.form.get("prodotto_id")
        vaschette = request.form.get("vaschette_prodotte", "").replace(",", ".")

        try:
            v = float(vaschette)
        except ValueError:
            flash("Numero di vaschette non valido.", "danger")
            return redirect(url_for("produzione"))

        if v <= 0:
            flash("Le vaschette devono essere maggiori di 0.", "danger")
            return redirect(url_for("produzione"))

        cur.execute(
            """
            INSERT INTO produzione (data, prodotto_id, vaschette_prodotte)
            VALUES (?, ?, ?)
            """,
            (data, prodotto_id, v),
        )
        conn.commit()
        conn.close()
        flash("Produzione registrata.", "success")
        return redirect(url_for("produzione"))

    # GET
    cur.execute("SELECT * FROM prodotti ORDER BY nome")
    prodotti = cur.fetchall()

    cur.execute(
        """
        SELECT pr.*,
               p.nome AS prodotto_nome,
               p.codice AS prodotto_codice,
               p.kg_per_vaschetta
        FROM produzione pr
        JOIN prodotti p ON p.id = pr.prodotto_id
        ORDER BY pr.data DESC, pr.id DESC
        """
    )
    rows = cur.fetchall()

    produzione_calc = []
    for r in rows:
        kg = r["vaschette_prodotte"] * r["kg_per_vaschetta"]
        produzione_calc.append(
            {
                "id": r["id"],
                "data": r["data"],
                "nome": r["prodotto_nome"],
                "codice": r["prodotto_codice"],
                "vaschette": r["vaschette_prodotte"],
                "kg": kg,
            }
        )

    conn.close()
    return render_template("produzione.html", prodotti=prodotti, produzione=produzione_calc)


# ---------------------- MAGAZZINO ----------------------


@app.route("/magazzino")
def magazzino():
    return render_template("magazzino.html", magazzino=calcola_magazzino())


# ---------------------- EXPORT LISTE CSV ----------------------


@app.route("/export/lista_carico")
def export_lista_carico():
    data_str = request.args.get("data")
    if not data_str:
        data_str = datetime.today().strftime("%Y-%m-%d")

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT o.data,
               c.nome AS cliente_nome,
               c.codice AS cliente_codice,
               p.nome AS prodotto_nome,
               p.codice AS prodotto_codice,
               p.kg_per_vaschetta,
               ro.qta_inserita,
               ro.tipo_qta
        FROM righe_ordine ro
        JOIN ordini o ON ro.ordine_id = o.id
        JOIN clienti c ON o.cliente_id = c.id
        JOIN prodotti p ON ro.prodotto_id = p.id
        WHERE o.data = ?
        ORDER BY c.nome, p.nome
        """,
        (data_str,),
    )
    rows = cur.fetchall()
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(
        ["Data", "Cliente", "Cod. Cliente", "Prodotto", "Cod. Prod.", "Vaschette", "Kg"]
    )

    for r in rows:
        q = r["qta_inserita"]
        t = r["tipo_qta"]
        kg_v = r["kg_per_vaschetta"]

        if t == "kg":
            kg = q
            vaschette = q / kg_v if kg_v else 0
        else:
            vaschette = q
            kg = q * kg_v

        writer.writerow(
            [
                r["data"],
                r["cliente_nome"],
                r["cliente_codice"] or "",
                r["prodotto_nome"],
                r["prodotto_codice"] or "",
                f"{vaschette:.2f}",
                f"{kg:.2f}",
            ]
        )

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"lista_carico_{data_str}.csv",
    )


@app.route("/export/magazzino")
def export_magazzino():
    mag = calcola_magazzino()

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(
        [
            "Cod",
            "Prodotto",
            "Kg/vaschetta",
            "Giacenza iniziale",
            "Prodotte",
            "Ordinate",
            "Giacenza vaschette",
            "Giacenza kg",
        ]
    )

    for r in mag:
        writer.writerow(
            [
                r["codice"] or "",
                r["nome"],
                f"{r['kg_per_vaschetta']:.3f}",
                f"{r['giacenza_iniziale_v']:.2f}",
                f"{r['prodotte_v']:.2f}",
                f"{r['ordinate_v']:.2f}",
                f"{r['giacenza_finale_v']:.2f}",
                f"{r['giacenza_finale_kg']:.2f}",
            ]
        )

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        mimetype="text/csv",
        as_attachment=True,
        download_name="magazzino.csv",
    )


# ---------------------- STAMPA CHECKLIST SINGOLO ORDINE ----------------------


@app.route("/ordini/<int:id>/stampa_checklist")
def stampa_checklist(id):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute(
        """
        SELECT o.id AS ordine_id,
               o.data,
               c.nome AS cliente_nome,
               c.codice AS cliente_codice
        FROM ordini o
        JOIN clienti c ON c.id = o.cliente_id
        WHERE o.id = ?
        """,
        (id,),
    )
    ordine = cur.fetchone()

    if ordine is None:
        conn.close()
        flash("Ordine non trovato.", "danger")
        return redirect(url_for("lista_ordini"))

    cur.execute(
        """
        SELECT ro.qta_inserita,
               ro.tipo_qta,
               p.nome AS prodotto_nome,
               p.codice AS prodotto_codice,
               p.kg_per_vaschetta
        FROM righe_ordine ro
        JOIN prodotti p ON p.id = ro.prodotto_id
        WHERE ro.ordine_id = ?
        ORDER BY p.nome
        """,
        (id,),
    )
    righe = cur.fetchall()
    conn.close()

    doc = Document()

    doc.add_heading("MAMMA CHE PASTA Srl - Checklist di Carico", level=1)

    doc.add_paragraph(f"Ordine n° {ordine['ordine_id']}  -  Data: {ordine['data']}")
    cliente_line = "Cliente: "
    if ordine["cliente_codice"]:
        cliente_line += f"[{ordine['cliente_codice']}] "
    cliente_line += ordine["cliente_nome"]
    doc.add_paragraph(cliente_line)
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Prodotto"
    hdr[1].text = "Kg"
    hdr[2].text = "Vaschette"
    hdr[3].text = "Check"

    tot_kg = 0
    tot_v = 0

    for r in righe:
        q = r["qta_inserita"]
        t = r["tipo_qta"]
        kg_v = r["kg_per_vaschetta"] or 0

        if t == "kg":
            kg = q
            vaschette = q / kg_v if kg_v else 0
        else:
            vaschette = q
            kg = q * kg_v

        tot_kg += kg
        tot_v += vaschette

        nome = r["prodotto_nome"]
        if r["prodotto_codice"]:
            nome = f"[{r['prodotto_codice']}] {nome}"

        row = table.add_row().cells
        row[0].text = nome
        row[1].text = f"{kg:.2f}"
        row[2].text = f"{vaschette:.2f}"
        row[3].text = "[ ]"

    doc.add_paragraph("")
    doc.add_paragraph(f"Totale kg: {tot_kg:.2f}")
    doc.add_paragraph(f"Totale vaschette: {tot_v:.2f}")
    doc.add_paragraph("")
    doc.add_paragraph("Firma magazziniere: ______________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Note:")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"checklist_{id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------- STAMPA ORDINI DEL GIORNO ----------------------


@app.route("/ordini/stampa_giorno")
def stampa_giorno():
    data_str = request.args.get("data")
    if not data_str:
        data_str = datetime.today().strftime("%Y-%m-%d")

    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute(
        """
        SELECT o.id AS ordine_id,
               o.data,
               c.nome AS cliente_nome,
               c.codice AS cliente_codice
        FROM ordini o
        JOIN clienti c ON o.cliente_id = c.id
        WHERE o.data = ?
        ORDER BY c.nome ASC, o.id ASC
        """,
        (data_str,),
    )
    ordini = cur.fetchall()

    if not ordini:
        conn.close()
        flash("Nessun ordine trovato per questa data.", "warning")
        return redirect(url_for("lista_ordini"))

    doc = Document()
    doc.add_heading(f"Ordini del giorno - {data_str}", level=1)

    for ordine in ordini:
        doc.add_paragraph("")
        doc.add_heading(
            f"Cliente: {ordine['cliente_nome']} ({ordine['cliente_codice'] or ''}) - Ordine n. {ordine['ordine_id']}",
            level=2,
        )

        cur.execute(
            """
            SELECT ro.qta_inserita,
                   ro.tipo_qta,
                   p.nome AS prodotto_nome,
                   p.codice AS prodotto_codice,
                   p.kg_per_vaschetta
            FROM righe_ordine ro
            JOIN prodotti p ON ro.prodotto_id = p.id
            WHERE ro.ordine_id = ?
            ORDER BY p.nome
            """,
            (ordine["ordine_id"],),
        )
        righe = cur.fetchall()

        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Prodotto"
        hdr[1].text = "Kg"
        hdr[2].text = "Vaschette"
        hdr[3].text = "Check"

        tot_kg = 0
        tot_v = 0

        for r in righe:
            q = r["qta_inserita"]
            t = r["tipo_qta"]
            kg_v = r["kg_per_vaschetta"] or 0

            if t == "kg":
                kg = q
                vaschette = q / kg_v if kg_v else 0
            else:
                vaschette = q
                kg = q * kg_v

            tot_kg += kg
            tot_v += vaschette

            nome = r["prodotto_nome"]
            if r["prodotto_codice"]:
                nome = f"[{r['prodotto_codice']}] {nome}"

            row = table.add_row().cells
            row[0].text = nome
            row[1].text = f"{kg:.2f}"
            row[2].text = f"{vaschette:.2f}"
            row[3].text = "[ ]"

        doc.add_paragraph(f"Totale Kg ordine: {tot_kg:.2f}")
        doc.add_paragraph(f"Totale vaschette ordine: {tot_v:.2f}")
        doc.add_paragraph("Firma magazziniere: _____________________________")

    conn.close()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"ordini_{data_str}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
@app.route("/statistiche")
def statistiche():
    conn = get_db_connection()
    cur = conn.cursor()

    # Prodotti più venduti
    cur.execute("""
        SELECT p.nome AS prodotto, SUM(r.qta_inserita) AS totale
        FROM righe_ordine r
        JOIN prodotti p ON p.id = r.prodotto_id
        GROUP BY p.id
        ORDER BY totale DESC
        LIMIT 10
    """)
    top_prodotti = cur.fetchall()

    # Clienti con il maggior numero di acquisti
    cur.execute("""
        SELECT c.nome AS cliente, SUM(r.qta_inserita) AS totale
        FROM righe_ordine r
        JOIN ordini o ON o.id = r.ordine_id
        JOIN clienti c ON c.id = o.cliente_id
        GROUP BY c.id
        ORDER BY totale DESC
        LIMIT 10
    """)
    top_clienti = cur.fetchall()

    # Andamento mensile
    cur.execute("""
        SELECT strftime('%Y-%m', o.data) AS mese,
               SUM(r.qta_inserita) AS totale
        FROM righe_ordine r
        JOIN ordini o ON o.id = r.ordine_id
        GROUP BY mese
        ORDER BY mese ASC
    """)
    andamento = cur.fetchall()

    conn.close()

    return render_template(
        "statistiche.html",
        top_prodotti=top_prodotti,
        top_clienti=top_clienti,
        andamento=andamento
    )



# ---------------------- MAIN ----------------------


if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=5000, debug=False)
